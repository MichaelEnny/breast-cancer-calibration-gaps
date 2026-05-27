"""
link_citations.py
-----------------
Adds internal hyperlinks between in-text citation numbers and their
corresponding entries in the References section of the manuscript docx.

Steps:
  1. Locate each reference paragraph and stamp a bookmark (ref1 ... ref19).
  2. Scan every body paragraph for superscript runs whose text is a plain
     integer 1-19, and wrap each in a w:hyperlink pointing to that bookmark.

Run:  python link_citations.py
Output: manuscript_draft_v3_linked.docx  (original file is not modified)
"""

import copy
import re
import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

INPUT   = "manuscript_draft_v3.docx"
OUTPUT  = "manuscript_draft_v3_linked.docx"
MAX_REF = 19


# ── helpers ──────────────────────────────────────────────────────────────────

def is_superscript(run):
    rPr = run._r.find(qn("w:rPr"))
    if rPr is not None:
        va = rPr.find(qn("w:vertAlign"))
        if va is not None and va.get(qn("w:val")) == "superscript":
            return True
    return False


def max_existing_bookmark_id(doc):
    """Return the highest w:bookmarkStart id already in the document."""
    highest = 0
    for bm in doc.element.iter(qn("w:bookmarkStart")):
        try:
            highest = max(highest, int(bm.get(qn("w:id"), 0)))
        except ValueError:
            pass
    return highest


def add_bookmark(para, name, bid):
    """Insert bookmarkStart / bookmarkEnd at the beginning of a paragraph."""
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bid))
    bm_start.set(qn("w:name"), name)

    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bid))

    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(bm_start)
        bm_start.addnext(bm_end)
    else:
        para._p.insert(0, bm_start)
        bm_start.addnext(bm_end)


def wrap_in_hyperlink(run, anchor):
    """Replace run with <w:hyperlink w:anchor='anchor'><w:r .../></w:hyperlink>."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)

    r_copy = copy.deepcopy(run._r)

    # Ensure rPr exists and carries the Hyperlink character style
    rPr = r_copy.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r_copy.insert(0, rPr)

    rStyle = rPr.find(qn("w:rStyle"))
    if rStyle is None:
        rStyle = OxmlElement("w:rStyle")
        rPr.insert(0, rStyle)
    rStyle.set(qn("w:val"), "Hyperlink")

    hyperlink.append(r_copy)

    run._r.addprevious(hyperlink)
    run._r.getparent().remove(run._r)
    return hyperlink


# ── main ─────────────────────────────────────────────────────────────────────

doc = Document(INPUT)

# ── Step 1: bookmark each reference paragraph ─────────────────────────────

ref_paras = []
in_refs = False
ref_entry_re = re.compile(r"^\d+\.\s")   # matches "1. ", "2. ", ... "19. "
for para in doc.paragraphs:
    if para.text.strip() == "References":
        in_refs = True
        continue
    if in_refs and para.text.strip():
        # Accept numbered entries ("1. Author...") as reference paragraphs.
        # Skip any other line (section heading after References, etc.).
        if ref_entry_re.match(para.text.strip()):
            ref_paras.append(para)

print(f"Found {len(ref_paras)} reference paragraph(s)")

next_bid = max_existing_bookmark_id(doc) + 1

for i, para in enumerate(ref_paras, 1):
    add_bookmark(para, f"ref{i}", next_bid)
    next_bid += 1
    print(f"  [ref{i}] → {para.text[:70]}")

# ── Step 2: hyperlink every superscript citation number ───────────────────

linked = 0
skipped = 0

ref_para_set = set(id(p) for p in ref_paras)
for para in doc.paragraphs:
    # Stop scanning body text once we hit the References heading
    if para.text.strip() == "References" and id(para) not in ref_para_set:
        break

    # Collect first so we don't mutate the list while iterating
    targets = []
    for run in para.runs:
        if is_superscript(run):
            text = run.text.strip()
            if re.match(r"^\d+$", text):
                num = int(text)
                if 1 <= num <= MAX_REF:
                    targets.append((run, num))
                else:
                    skipped += 1

    for run, num in targets:
        wrap_in_hyperlink(run, f"ref{num}")
        linked += 1

print(f"\nLinked {linked} citation(s)  |  skipped {skipped} out-of-range superscript(s)")

doc.save(OUTPUT)
print(f"Saved → {OUTPUT}")
